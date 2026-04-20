#!/usr/bin/env python3
"""
Document Translator — translates DOCX, PDF, and EPUB files offline using Argos Translate.

Similar to using Whisper offline for transcription, this script downloads
neural translation models once and runs everything locally — no API costs.

Usage:
    python translate.py document.docx
    python translate.py document.pdf --source en --target es
    python translate.py book.epub --target pt
    python translate.py --setup en pt        # download English→Portuguese model
    python translate.py --list-langs         # show available languages
"""

import argparse
import html
import logging
import os
import re
import sys
import textwrap
import time
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString
from tqdm import tqdm

logging.basicConfig(
    level=logging.ERROR,
    format="%(message)s",
)
logger = logging.getLogger(__name__)

import argostranslate.utils
argostranslate.utils.logger.setLevel(logging.ERROR)
for _h in argostranslate.utils.logger.handlers[:]:
    argostranslate.utils.logger.removeHandler(_h)
logging.getLogger("stanza").setLevel(logging.CRITICAL)

DEFAULT_SOURCE_LANG = "en"
DEFAULT_TARGET_LANG = "pb"
DEFAULT_CHUNK_SIZE = 2000  # characters per translation call

CLEAN_FORMAT = {
    "font_name": "Calibri",
    "font_size": 11,
    "line_spacing": 1.15,
    "space_after_pt": 6,
    "margin_cm": 2.54,
}


# ---------------------------------------------------------------------------
# Argos Translate helpers
# ---------------------------------------------------------------------------

def ensure_model_installed(from_code: str, to_code: str) -> None:
    """Download and install the language pair model if not already present."""
    import argostranslate.package
    import argostranslate.translate

    installed = argostranslate.translate.get_installed_languages()
    from_lang = next((l for l in installed if l.code == from_code), None)

    if from_lang:
        translation = next(
            (t for t in from_lang.translations_to if t.to_lang.code == to_code),
            None,
        )
        if translation:
            logger.info("Model %s→%s already installed.", from_code, to_code)
            return

    print(f"Downloading model {from_code}→{to_code} (~100 MB, first time only)...")
    argostranslate.package.update_package_index()
    available = argostranslate.package.get_available_packages()
    pkg = next(
        (p for p in available if p.from_code == from_code and p.to_code == to_code),
        None,
    )
    if pkg is None:
        print(f"Error: no translation model found for {from_code}→{to_code}. "
              "Run --list-langs to see available languages.")
        sys.exit(1)

    path = pkg.download()
    argostranslate.package.install_from_path(path)
    print(f"Model {from_code}→{to_code} installed.")


def list_available_languages() -> None:
    """Print all available language pairs from the Argos package index."""
    import argostranslate.package

    argostranslate.package.update_package_index()
    available = argostranslate.package.get_available_packages()

    print("\nAvailable language pairs:\n")
    print(f"  {'From':<25} {'Code':<8} {'To':<25} {'Code':<8}")
    print(f"  {'-'*25} {'-'*8} {'-'*25} {'-'*8}")
    for pkg in sorted(available, key=lambda p: (p.from_code, p.to_code)):
        print(f"  {pkg.from_name:<25} {pkg.from_code:<8} {pkg.to_name:<25} {pkg.to_code:<8}")

    installed = argostranslate.package.get_installed_packages()
    if installed:
        print(f"\nInstalled models ({len(installed)}):")
        for pkg in installed:
            print(f"  {pkg.from_name} ({pkg.from_code}) → {pkg.to_name} ({pkg.to_code})")
    else:
        print("\nNo models installed yet. Use --setup <from> <to> to install one.")
    print()


def get_language_names(from_code: str, to_code: str) -> tuple[str, str]:
    """Resolve language codes to human-readable names."""
    import argostranslate.translate

    installed = argostranslate.translate.get_installed_languages()
    from_name = next((l.name for l in installed if l.code == from_code), from_code)
    to_name = next((l.name for l in installed if l.code == to_code), to_code)
    return from_name, to_name


def translate_text_offline(text: str, from_code: str, to_code: str) -> str:
    """Translate text using a locally installed Argos model."""
    import argostranslate.translate

    logging.getLogger("stanza").setLevel(logging.CRITICAL)
    return argostranslate.translate.translate(text, from_code, to_code)


# ---------------------------------------------------------------------------
# Document readers
# ---------------------------------------------------------------------------

def read_docx(path: str) -> list[dict]:
    """Return a list of paragraph dicts with text and style metadata."""
    import docx

    doc = docx.Document(path)
    paragraphs: list[dict] = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append({
                "text": para.text,
                "style": para.style.name if para.style else None,
                "runs": [
                    {
                        "text": run.text,
                        "bold": run.bold,
                        "italic": run.italic,
                        "underline": run.underline,
                        "font_name": run.font.name,
                        "font_size": run.font.size,
                    }
                    for run in para.runs
                ],
            })
    return paragraphs


def read_pdf(path: str) -> list[dict]:
    """Extract text from PDF pages, stripping repeated headers/footers."""
    import pdfplumber
    from collections import Counter

    raw_pages: list[dict] = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text and text.strip():
                raw_pages.append({"page": i + 1, "text": text})

    if len(raw_pages) < 3:
        return raw_pages

    line_counts: Counter[str] = Counter()
    for p in raw_pages:
        lines = p["text"].split("\n")
        first_last = set()
        for line in lines[:2] + lines[-2:]:
            stripped = line.strip()
            if stripped:
                first_last.add(stripped)
        for line in first_last:
            line_counts[line] += 1

    threshold = len(raw_pages) * 0.6
    repeated = {line for line, count in line_counts.items() if count >= threshold}

    pages: list[dict] = []
    for p in raw_pages:
        filtered = "\n".join(
            line for line in p["text"].split("\n")
            if line.strip() not in repeated
        )
        if filtered.strip():
            pages.append({"page": p["page"], "text": filtered})

    return pages


def read_epub(path: str) -> list[dict]:
    """Extract HTML content from EPUB chapters."""
    import ebooklib
    from ebooklib import epub

    book = epub.read_epub(path)
    chapters: list[dict] = []
    for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
        content = item.get_content().decode("utf-8", errors="replace")
        soup = BeautifulSoup(content, "lxml")
        body = soup.find("body")
        text = body.get_text(separator="\n", strip=True) if body else soup.get_text(separator="\n", strip=True)
        if text.strip():
            chapters.append({
                "id": item.get_id(),
                "name": item.get_name(),
                "html": content,
                "text": text,
            })
    return chapters


# ---------------------------------------------------------------------------
# Paragraph reconstruction
# ---------------------------------------------------------------------------

_SENTENCE_END = re.compile(r"[.!?:;]\s*$")
_PAGE_NUM = re.compile(r"^\s*\d{1,4}\s*$")
_HEADING_LIKE = re.compile(
    r"^(\d{1,2}[\s.)\-–—]+[A-ZÀ-Ú]|Cap[ií]tulo|CAPÍTULO|Introdu|Conclus|Refer[êe]ncias)",
    re.IGNORECASE,
)


def _fix_pdf_ligatures(text: str) -> str:
    """Fix broken typographic ligatures from PDF extraction.

    PDF extractors sometimes split ligature glyphs ("fi", "fl", "ff")
    with a space: "fi nanceiro" → "financeiro", "refl exão" → "reflexão".
    Also replaces Unicode ligature codepoints (U+FB00–FB04).
    """
    text = text.replace("\ufb00", "ff")
    text = text.replace("\ufb01", "fi")
    text = text.replace("\ufb02", "fl")
    text = text.replace("\ufb03", "ffi")
    text = text.replace("\ufb04", "ffl")

    _lower = r"[a-záàâãéèêíïóôõúüç]"
    text = re.sub(rf"fi\s+(?={_lower})", "fi", text)
    text = re.sub(rf"fl\s+(?={_lower})", "fl", text)
    text = re.sub(rf"ff\s+(?={_lower})", "ff", text)
    return text


def merge_lines_into_paragraphs(text: str) -> str:
    """Reconstruct proper paragraphs from PDF-extracted text.

    PDF extraction produces one line per visual line, breaking paragraphs
    arbitrarily.  This merges consecutive lines that belong to the same
    paragraph.

    Key rules (aligned with doc2audio's ``_rebuild_pdf_paragraphs``):
    - Blank lines only flush a paragraph when the accumulated text ends with
      sentence-terminal punctuation (``.!?:;``).  Otherwise the blank line is
      treated as extraction noise and the paragraph continues.
    - Lines that look like page numbers are discarded.
    - Lines that look like headings (ALL CAPS or "Capítulo …" patterns) always
      start a new paragraph.
    - Bullet / numbered-list items always start a new paragraph.
    - A new paragraph starts when the previous line ends with sentence-terminal
      punctuation **and** the next line begins with an uppercase letter.
    """
    lines = text.splitlines()
    if not lines:
        return text

    paragraphs: list[str] = []
    current: list[str] = []

    for line in lines:
        stripped = line.strip()

        if not stripped:
            if current and _SENTENCE_END.search(current[-1]):
                paragraphs.append(" ".join(current))
                current = []
            continue

        if _PAGE_NUM.match(stripped):
            continue

        is_heading = (
            stripped.isupper() and len(stripped) > 3
        ) or _HEADING_LIKE.match(stripped)

        is_list_item = (
            stripped.startswith(("•", "-", "–", "—", "▪"))
        ) or (
            len(stripped) > 2 and stripped[0].isdigit() and stripped[1] in ".)"
        )

        if is_heading or is_list_item:
            if current:
                paragraphs.append(" ".join(current))
                current = []
            if is_heading:
                paragraphs.append(stripped)
            else:
                current.append(stripped)
            continue

        if not current:
            current.append(stripped)
            continue

        prev = current[-1]
        starts_upper = stripped[0].isupper() if stripped else False

        if _SENTENCE_END.search(prev) and starts_upper:
            paragraphs.append(" ".join(current))
            current = [stripped]
        else:
            current.append(stripped)

    if current:
        paragraphs.append(" ".join(current))

    result = "\n".join(paragraphs)
    return _fix_pdf_ligatures(result)


# ---------------------------------------------------------------------------
# Text chunking and batch translation
# ---------------------------------------------------------------------------

def chunk_text(text: str, max_chars: int = DEFAULT_CHUNK_SIZE) -> list[str]:
    """Split text into chunks at paragraph boundaries respecting max_chars."""
    paragraphs = text.split("\n")
    chunks: list[str] = []
    current: list[str] = []
    current_len = 0

    for para in paragraphs:
        para_len = len(para) + 1
        if current_len + para_len > max_chars and current:
            chunks.append("\n".join(current))
            current = []
            current_len = 0
        current.append(para)
        current_len += para_len

    if current:
        chunks.append("\n".join(current))

    return chunks


def translate_chunks(
    chunks: list[str],
    from_code: str,
    to_code: str,
    label: str = "",
) -> list[str]:
    """Translate a list of text chunks with progress bar."""
    translated: list[str] = []
    desc = f"Translating {label}" if label else "Translating"
    for chunk in tqdm(chunks, desc=desc, unit="chunk"):
        result = translate_text_offline(chunk, from_code, to_code)
        translated.append(result)
    return translated


# ---------------------------------------------------------------------------
# Document writers
# ---------------------------------------------------------------------------

def write_docx(paragraphs: list[dict], translated_texts: list[str], output_path: str) -> None:
    """Create a new DOCX with translated content preserving basic styles."""
    import docx

    doc = docx.Document()

    for para_data, translated in zip(paragraphs, translated_texts):
        p = doc.add_paragraph()
        style_name = para_data.get("style")
        if style_name:
            try:
                p.style = style_name
            except KeyError:
                pass

        original_runs = para_data.get("runs", [])
        if len(original_runs) <= 1:
            run = p.add_run(translated)
            if original_runs:
                fmt = original_runs[0]
                run.bold = fmt.get("bold")
                run.italic = fmt.get("italic")
                run.underline = fmt.get("underline")
                if fmt.get("font_name"):
                    run.font.name = fmt["font_name"]
                if fmt.get("font_size"):
                    run.font.size = fmt["font_size"]
        else:
            run = p.add_run(translated)
            fmt = original_runs[0]
            run.bold = fmt.get("bold")
            run.italic = fmt.get("italic")

    doc.save(output_path)
    logger.info("Saved translated DOCX: %s", output_path)


def write_pdf_as_docx(
    pages: list[dict],
    translated_texts: list[str],
    output_path: str,
    title: str = "",
) -> None:
    """
    Write translated PDF content as DOCX.
    Reconstructs paragraphs and adds page separators.
    """
    import docx

    doc = docx.Document()

    if title:
        doc.add_paragraph(title)
        doc.add_paragraph("")

    for page_data, translated in zip(pages, translated_texts):
        doc.add_paragraph(f"— p. {page_data['page']} —")
        merged = merge_lines_into_paragraphs(translated)
        for para_text in merged.split("\n"):
            if para_text.strip():
                doc.add_paragraph(para_text)

    doc.save(output_path)


def write_epub(
    original_path: str,
    chapters: list[dict],
    translated_texts: list[str],
    output_path: str,
) -> None:
    """Create a translated EPUB preserving structure and metadata."""
    import ebooklib
    from ebooklib import epub

    book = epub.read_epub(original_path)

    translated_map: dict[str, str] = {}
    for chapter, translated in zip(chapters, translated_texts):
        translated_map[chapter["id"]] = translated

    for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
        item_id = item.get_id()
        if item_id not in translated_map:
            continue

        html_content = item.get_content().decode("utf-8", errors="replace")
        soup = BeautifulSoup(html_content, "lxml")
        body = soup.find("body")
        target = body if body else soup

        translated_paragraphs = translated_map[item_id].split("\n")
        _replace_text_nodes(target, translated_paragraphs)

        item.set_content(str(soup).encode("utf-8"))

    epub.write_epub(output_path, book)
    logger.info("Saved translated EPUB: %s", output_path)


def _setup_pdf_font(pdf) -> str:
    """Register a Unicode TTF font and return its family name."""
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/TTF/DejaVuSans.ttf",
        "C:\\Windows\\Fonts\\arial.ttf",
    ]
    for path in candidates:
        if os.path.isfile(path):
            pdf.add_font("UniFont", "", path)
            pdf.add_font("UniFont", "B", path)
            return "UniFont"
    return "Helvetica"


def write_pdf(pages: list[dict], translated_texts: list[str], output_path: str, title: str = "") -> None:
    """Generate a formatted PDF from translated text using fpdf2."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=25)
    font = _setup_pdf_font(pdf)

    pdf.add_page()
    pdf.set_font(font, size=14, style="B")
    if title:
        pdf.multi_cell(w=0, text=title, align="C")
        pdf.ln(10)

    for page_data, translated in zip(pages, translated_texts):
        merged = merge_lines_into_paragraphs(translated)

        for para_text in merged.split("\n"):
            stripped = para_text.strip()
            if not stripped:
                continue
            pdf.set_font(font, size=10)
            pdf.set_text_color(0, 0, 0)
            pdf.set_x(pdf.l_margin + 10)
            pdf.multi_cell(w=0, text=stripped, align="J")
            pdf.ln(2)

        pdf.set_font(font, size=7)
        pdf.set_text_color(150, 150, 150)
        pdf.cell(w=0, text=f"-- p. {page_data['page']} --", align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(0, 0, 0)
        pdf.ln(4)

    pdf.output(output_path)


def write_html(paragraphs: list[str], output_path: str, title: str = "") -> None:
    """Create a simple standalone HTML document from translated paragraphs."""
    safe_title = html.escape(title or "Translated document")
    html_lines = [
        "<!doctype html>",
        '<html lang="en">',
        "<head>",
        '  <meta charset="utf-8">',
        '  <meta name="viewport" content="width=device-width, initial-scale=1">',
        f"  <title>{safe_title}</title>",
        "  <style>",
        "    body { font-family: Arial, sans-serif; max-width: 840px; margin: 40px auto; line-height: 1.6; color: #111; padding: 0 16px; }",
        "    h1 { font-size: 1.6rem; margin-bottom: 1rem; }",
        "    p { margin: 0 0 0.9rem; text-align: justify; }",
        "    .page-sep { color: #666; text-align: center; font-size: 0.85rem; margin: 1.2rem 0; }",
        "  </style>",
        "</head>",
        "<body>",
        f"  <h1>{safe_title}</h1>",
    ]

    for para in paragraphs:
        stripped = para.strip()
        if not stripped:
            continue
        safe_para = html.escape(stripped)
        if safe_para.startswith("— p.") and safe_para.endswith("—"):
            html_lines.append(f'  <div class="page-sep">{safe_para}</div>')
        else:
            html_lines.append(f"  <p>{safe_para}</p>")

    html_lines.extend(["</body>", "</html>"])
    Path(output_path).write_text("\n".join(html_lines), encoding="utf-8")


def format_docx(path: str) -> None:
    """Apply clean, uniform formatting to an existing DOCX file."""
    import docx
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH

    fmt = CLEAN_FORMAT
    doc = docx.Document(path)

    for section in doc.sections:
        section.top_margin = Cm(fmt["margin_cm"])
        section.bottom_margin = Cm(fmt["margin_cm"])
        section.left_margin = Cm(fmt["margin_cm"])
        section.right_margin = Cm(fmt["margin_cm"])

    is_first_text = True
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        pf = para.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = fmt["line_spacing"]
        pf.space_after = Pt(fmt["space_after_pt"])

        is_page_sep = text.startswith("— p.") and text.endswith("—")

        if is_first_text and not is_page_sep:
            run = para.clear().add_run(text)
            run.font.name = fmt["font_name"]
            run.font.size = Pt(16)
            run.bold = True
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.space_after = Pt(18)
            is_first_text = False
            continue

        if is_page_sep:
            run = para.clear().add_run(text)
            run.font.name = fmt["font_name"]
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(150, 150, 150)
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.space_before = Pt(12)
            pf.space_after = Pt(12)
            continue

        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.first_line_indent = Cm(1.25)

        if para.runs:
            for run in para.runs:
                run.font.name = fmt["font_name"]
                run.font.size = Pt(fmt["font_size"])
                run.bold = False
                run.italic = False
        else:
            run = para.clear().add_run(text)
            run.font.name = fmt["font_name"]
            run.font.size = Pt(fmt["font_size"])

        is_first_text = False

    doc.save(path)


def _replace_text_nodes(element, translated_lines: list[str]) -> None:
    """Walk the DOM tree and replace text nodes with translated content."""
    line_iter = iter(translated_lines)

    def _walk(node):
        if isinstance(node, NavigableString) and node.string and node.string.strip():
            try:
                replacement = next(line_iter)
                node.replace_with(replacement)
            except StopIteration:
                pass
        elif hasattr(node, "children"):
            for child in list(node.children):
                _walk(child)

    _walk(element)


# ---------------------------------------------------------------------------
# Orchestration
# ---------------------------------------------------------------------------

def get_output_path(input_path: str, to_code: str, ext_override: str | None = None) -> str:
    """Generate output file path with language prefix."""
    p = Path(input_path)
    ext = ext_override if ext_override else p.suffix
    return str(p.with_name(f"{to_code} - {p.stem}{ext}"))


def _is_within_directory(candidate_path: str, base_dir: str) -> bool:
    """Return True when candidate_path resolves under base_dir."""
    try:
        return os.path.commonpath([candidate_path, base_dir]) == base_dir
    except ValueError:
        return False


def process_docx(input_path: str, from_code: str, to_code: str, chunk_size: int) -> tuple[str, str]:
    logger.info("Reading DOCX: %s", input_path)
    paragraphs = read_docx(input_path)
    logger.info("Found %d paragraphs", len(paragraphs))

    texts = [p["text"] for p in paragraphs]
    sep = "\n|||SEP|||\n"
    chunks = chunk_text(sep.join(texts), chunk_size)
    raw_translated = translate_chunks(chunks, from_code, to_code, label=Path(input_path).name)
    full_translated = sep.join(raw_translated)
    translated = [t.strip() for t in full_translated.split("|||SEP|||")]

    if len(translated) < len(paragraphs):
        translated.extend([""] * (len(paragraphs) - len(translated)))
    translated = translated[: len(paragraphs)]

    output_path = get_output_path(input_path, to_code)
    write_docx(paragraphs, translated, output_path)
    html_path = get_output_path(input_path, to_code, ext_override=".html")
    title = Path(input_path).stem
    write_html(translated, html_path, title=title)
    return output_path, html_path


def process_pdf(input_path: str, from_code: str, to_code: str, chunk_size: int) -> tuple[str, str, str]:
    logger.info("Reading PDF: %s", input_path)
    pages = read_pdf(input_path)
    logger.info("Found %d pages with text", len(pages))

    translated: list[str] = []
    for page in tqdm(pages, desc="Translating pages", unit="page"):
        chunks = chunk_text(page["text"], chunk_size)
        page_parts = [translate_text_offline(c, from_code, to_code) for c in chunks]
        translated.append("\n".join(page_parts))

    title = Path(input_path).stem

    docx_path = get_output_path(input_path, to_code, ext_override=".docx")
    write_pdf_as_docx(pages, translated, docx_path, title=title)

    pdf_path = get_output_path(input_path, to_code, ext_override=".pdf")
    write_pdf(pages, translated, pdf_path, title=title)

    html_path = get_output_path(input_path, to_code, ext_override=".html")
    html_paragraphs: list[str] = []
    for page_data, translated_page in zip(pages, translated):
        html_paragraphs.append(f"— p. {page_data['page']} —")
        merged = merge_lines_into_paragraphs(translated_page)
        html_paragraphs.extend(merged.split("\n"))
    write_html(html_paragraphs, html_path, title=title)

    return docx_path, pdf_path, html_path


def process_epub(input_path: str, from_code: str, to_code: str, chunk_size: int) -> str:
    logger.info("Reading EPUB: %s", input_path)
    chapters = read_epub(input_path)
    logger.info("Found %d chapters", len(chapters))

    translated: list[str] = []
    for ch in tqdm(chapters, desc="Translating chapters", unit="ch"):
        chunks = chunk_text(ch["text"], chunk_size)
        ch_parts = [translate_text_offline(c, from_code, to_code) for c in chunks]
        translated.append("\n".join(ch_parts))

    output_path = get_output_path(input_path, to_code)
    write_epub(input_path, chapters, translated, output_path)
    return output_path


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Translate documents (DOCX, PDF, EPUB) offline using Argos Translate.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=textwrap.dedent("""\
            Examples:
              python translate.py document.docx                          # EN → PT (default)
              python translate.py paper.pdf --source en --target es      # EN → ES
              python translate.py novel.epub --source fr                 # FR → PT
              python translate.py --setup en pt                          # download model
              python translate.py --list-langs                           # show languages

            The first run for a language pair downloads the model (~100 MB).
            After that, everything runs 100%% offline — no API, no cost.
        """),
    )

    parser.add_argument("input", nargs="?", help="Path to the document to translate")
    parser.add_argument(
        "--source", "-s",
        default=DEFAULT_SOURCE_LANG,
        help=f"Source language code (default: {DEFAULT_SOURCE_LANG})",
    )
    parser.add_argument(
        "--target", "-t",
        default=DEFAULT_TARGET_LANG,
        help=f"Target language code (default: {DEFAULT_TARGET_LANG})",
    )
    parser.add_argument(
        "--chunk-size",
        type=int,
        default=DEFAULT_CHUNK_SIZE,
        help=f"Max characters per translation chunk (default: {DEFAULT_CHUNK_SIZE})",
    )
    parser.add_argument(
        "--output", "-o",
        default=None,
        help="Custom output file path (auto-generated if omitted)",
    )
    parser.add_argument(
        "--no-format",
        action="store_true",
        help="Skip clean formatting on output DOCX (Calibri 11pt, 1.15 spacing applied by default)",
    )
    parser.add_argument(
        "--setup",
        nargs=2,
        metavar=("FROM", "TO"),
        help="Download a translation model, e.g. --setup en pt",
    )
    parser.add_argument(
        "--list-langs",
        action="store_true",
        help="List all available language pairs and exit",
    )
    parser.add_argument(
        "--verbose",
        action="store_true",
        help="Enable debug logging",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()

    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)

    if args.list_langs:
        list_available_languages()
        return

    if args.setup:
        from_code, to_code = args.setup
        ensure_model_installed(from_code, to_code)
        return

    if not args.input:
        logger.error("No input file provided. Use --help for usage information.")
        sys.exit(1)

    input_path = os.path.realpath(args.input)
    if not Path(input_path).is_file():
        logger.error("File not found: %s", input_path)
        sys.exit(1)

    suffix = Path(input_path).suffix.lower()
    supported = {".docx", ".pdf", ".epub"}
    if suffix not in supported:
        logger.error("Unsupported format '%s'. Supported: %s", suffix, ", ".join(sorted(supported)))
        sys.exit(1)

    from_code = args.source
    to_code = args.target

    ensure_model_installed(from_code, to_code)

    from_name, to_name = get_language_names(from_code, to_code)
    msg = f"Translating from {from_name} to {to_name}"
    print(f"\n── {msg} ──\n")

    processors = {
        ".docx": process_docx,
        ".pdf": process_pdf,
        ".epub": process_epub,
    }

    start_time = time.time()

    result = processors[suffix](input_path, from_code, to_code, args.chunk_size)

    if suffix == ".pdf":
        docx_path, pdf_path, html_path = result
        if not args.no_format:
            format_docx(docx_path)
        output_files = [docx_path, pdf_path, html_path]
    elif suffix == ".docx":
        docx_path, html_path = result
        output_files = [docx_path, html_path]
    else:
        output_files = [result]

    if args.output:
        cwd = os.path.realpath(os.getcwd())

        output_name = Path(args.output).name
        if output_name != args.output:
            print("Error: --output must be a file name, without directory components.")
            sys.exit(1)

        resolved_dest = os.path.realpath(os.path.join(cwd, output_name))
        if not _is_within_directory(resolved_dest, cwd):
            print("Error: output path must stay inside the current working directory.")
            sys.exit(1)

        resolved_source = os.path.realpath(output_files[0])
        if not _is_within_directory(resolved_source, cwd):
            print("Error: generated output path is outside the working directory.")
            sys.exit(1)

        os.replace(resolved_source, resolved_dest)
        output_files[0] = resolved_dest

    elapsed = time.time() - start_time
    minutes, seconds = divmod(int(elapsed), 60)
    time_str = f"{minutes}m{seconds:02d}s" if minutes else f"{seconds}s"

    print(f"\nDone in {time_str} — {from_code}→{to_code}")
    for f in output_files:
        print(f"  {f}")


if __name__ == "__main__":
    main()
